"""
Resume Tailoring Module
==========================
Builds a "tailored" version of a candidate's resume by incorporating
user-approved missing skills, without inventing any claims the user didn't
explicitly provide. Two paths:

1. DOCX upload: edit the real python-docx Document object directly, so the
   original formatting/styling is preserved as much as possible.
2. PDF upload: we can't safely edit a PDF's absolute-positioned layout, so we
   rebuild a clean, simple DOCX from the extracted text instead.

In both cases, every word in the output is either:
  (a) already present in the original resume, or
  (b) text the user explicitly typed in themselves for a specific skill.
Nothing is auto-generated and silently inserted as a factual claim.
"""

import re
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SECTION_HEADING_PATTERNS = [
    r"^skills?$", r"^technical skills?$", r"^core skills?$",
    r"^key skills?$", r"^skills? (and|&) (expertise|competencies)$",
]
_HEADING_RE = re.compile("|".join(SECTION_HEADING_PATTERNS), re.IGNORECASE)


def _is_heading_paragraph(paragraph):
    """
    Heuristic: a paragraph counts as a heading if its style name contains
    'Heading', OR its text is short, fully capitalized/title-cased, and has
    no trailing punctuation -- covers resumes that fake headings with bold
    text instead of real Word heading styles.
    """
    if "Heading" in paragraph.style.name:
        return True
    text = paragraph.text.strip()
    if not text or len(text) > 40:
        return False
    if text.isupper():
        return True
    return False


def find_skills_section_paragraph(doc):
    """
    Returns a tuple (heading_index, content_index) where content_index is the
    paragraph index holding the skills *content* (the first non-empty
    paragraph after the heading that is NOT itself another section heading),
    or None for content_index if the heading has no content before the next
    section starts (or before the document ends). Returns None entirely if
    no skills-section heading was found at all.
    """
    for i, para in enumerate(doc.paragraphs):
        if _HEADING_RE.match(para.text.strip()):
            for j in range(i + 1, len(doc.paragraphs)):
                candidate = doc.paragraphs[j]
                if not candidate.text.strip():
                    continue
                if _is_heading_paragraph(candidate):
                    # Hit the next section before finding any skills content.
                    return (i, None)
                return (i, j)
            return (i, None)
    return None


def add_skills_to_docx(doc, new_skills):
    """
    Appends new_skills to the existing skills section content line if one is
    found; if a skills heading exists but has no content line yet, inserts
    one right after it; otherwise creates a new 'SKILLS' section at the end
    of the document. Only ever adds skill names the user explicitly chose to
    include -- never invents context around them.
    """
    if not new_skills:
        return doc

    section = find_skills_section_paragraph(doc)

    if section is not None:
        heading_idx, content_idx = section

        if content_idx is not None:
            para = doc.paragraphs[content_idx]
            existing_text = para.text.rstrip()
            separator = ", " if existing_text and not existing_text.endswith(",") else " "
            addition = ", ".join(new_skills)
            run = para.add_run(f"{separator}{addition}")
            run.font.size = para.runs[0].font.size if para.runs else Pt(11)
        else:
            # Heading exists but has no content paragraph after it.
            heading_para = doc.paragraphs[heading_idx]
            if heading_idx == len(doc.paragraphs) - 1:
                # Heading is the very last paragraph -- just append at the end.
                doc.add_paragraph(", ".join(new_skills))
            else:
                # Insert a new content paragraph directly after the heading by
                # inserting it before whatever paragraph currently follows the heading.
                following_para = doc.paragraphs[heading_idx + 1]
                new_para = following_para.insert_paragraph_before(", ".join(new_skills))
    else:
        doc.add_paragraph()  # spacing
        heading = doc.add_paragraph()
        heading_run = heading.add_run("SKILLS")
        heading_run.bold = True
        heading_run.font.size = Pt(12)

        doc.add_paragraph(", ".join(new_skills))

    return doc


def add_bullets_to_docx(doc, bullets):
    """
    Appends a new 'ADDITIONAL EXPERIENCE' section at the end of the document
    containing only bullets the user explicitly wrote themselves.
    bullets: list of dicts {"skill": str, "text": str}
    """
    if not bullets:
        return doc

    doc.add_paragraph()  # spacing
    heading = doc.add_paragraph()
    heading_run = heading.add_run("ADDITIONAL EXPERIENCE")
    heading_run.bold = True
    heading_run.font.size = Pt(12)

    for b in bullets:
        bullet_para = doc.add_paragraph(style="List Bullet")
        bullet_para.add_run(b["text"])

    return doc


def build_tailored_docx_from_existing(original_docx_path, skills_only, bullets, output_path):
    """
    For DOCX uploads: load the real document and edit it in place.
    skills_only: list of skill name strings to add to the Skills section
    bullets: list of {"skill", "text"} dicts to add as new bullets
    """
    doc = docx.Document(original_docx_path)
    add_skills_to_docx(doc, skills_only)
    add_bullets_to_docx(doc, bullets)
    doc.save(output_path)
    return output_path


def build_tailored_docx_from_text(original_text, skills_only, bullets, output_path):
    """
    For PDF uploads: we can't safely edit the original PDF's layout, so we
    rebuild a clean, simple DOCX containing the original extracted text as
    plain paragraphs, then append the same Skills/Bullet additions.
    This intentionally does not try to recreate the original's visual
    design -- only its actual textual content.
    """
    doc = docx.Document()

    title = doc.add_paragraph()
    title_run = title.add_run("Resume")
    title_run.bold = True
    title_run.font.size = Pt(14)
    note = doc.add_paragraph()
    note_run = note.add_run(
        "(Rebuilt from the uploaded PDF's text content -- original visual "
        "formatting could not be preserved.)"
    )
    note_run.italic = True
    note_run.font.size = Pt(8)
    doc.add_paragraph()

    for line in original_text.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    add_skills_to_docx(doc, skills_only)
    add_bullets_to_docx(doc, bullets)
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    import shutil
    shutil.copy("/tmp/test_resume.docx", "/tmp/test_resume_copy.docx")

    build_tailored_docx_from_existing(
        "/tmp/test_resume_copy.docx",
        skills_only=["AWS", "Docker"],
        bullets=[{"skill": "Statistics",
                  "text": "Applied statistical hypothesis testing to validate A/B test results across three product launches."}],
        output_path="/tmp/tailored_test.docx",
    )

    doc = docx.Document("/tmp/tailored_test.docx")
    for p in doc.paragraphs:
        print(repr(p.text))
